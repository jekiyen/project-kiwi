import io

from docx import Document
from pypdf import PdfReader


class TextExtractionError(Exception):
    """Raised when a file's text can't be extracted (corrupted, encrypted, unsupported, empty)."""


def extract_text_from_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise TextExtractionError(f"PDF is password-protected: {exc}") from exc
        text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except TextExtractionError:
        raise
    except Exception as exc:
        raise TextExtractionError(f"Failed to read PDF: {exc}") from exc

    if not text:
        raise TextExtractionError(
            "No extractable text found in this PDF — it may be a scanned image with no text layer."
        )
    return text


def extract_text_from_docx(data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        text = "\n".join(parts).strip()
    except TextExtractionError:
        raise
    except Exception as exc:
        raise TextExtractionError(f"Failed to read DOCX: {exc}") from exc

    if not text:
        raise TextExtractionError("No extractable text found in this DOCX file.")
    return text


def extract_text(data: bytes, file_type: str) -> str:
    if file_type == "pdf":
        return extract_text_from_pdf(data)
    if file_type == "docx":
        return extract_text_from_docx(data)
    raise TextExtractionError(f"Unsupported file type: {file_type!r}")
